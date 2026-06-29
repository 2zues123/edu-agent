#!/usr/bin/env python3
"""Build retrieval chunks from raw knowledge-base documents."""

from __future__ import annotations

import argparse
import hashlib
import json
import re
import shutil
import subprocess
from dataclasses import asdict, dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Iterable


RAW_DIR = Path("data/raw")
OUTPUT_FILE = Path("data/processed/chunks.jsonl")
REPORT_FILE = Path("data/processed/build_report.json")

SUPPORTED_SUFFIXES = {".doc", ".docx", ".pdf", ".txt", ".md"}


@dataclass
class SourceDocument:
    path: Path
    category: str
    title: str
    suffix: str
    doc_id: str
    source_url: str = ""
    site: str = ""
    published_at: str = ""


@dataclass
class Chunk:
    chunk_id: str
    doc_id: str
    title: str
    category: str
    source_file: str
    chunk_index: int
    heading: str
    text: str
    source_url: str = ""
    site: str = ""
    published_at: str = ""


def stable_id(value: str, length: int = 16) -> str:
    return hashlib.sha1(value.encode("utf-8")).hexdigest()[:length]


def normalize_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = text.replace("\u2028", "\n").replace("\u2029", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def infer_title(path: Path) -> str:
    return path.stem.strip()


def infer_category(path: Path, raw_dir: Path) -> str:
    relative_parts = path.relative_to(raw_dir).parts
    if relative_parts and relative_parts[0] in {"web", "web_attachments"}:
        return "web"
    return path.parent.name


def parse_front_matter(text: str) -> tuple[dict[str, str], str]:
    if not text.startswith("---\n"):
        return {}, text
    end = text.find("\n---", 4)
    if end < 0:
        return {}, text

    metadata: dict[str, str] = {}
    raw_metadata = text[4:end].splitlines()
    for line in raw_metadata:
        if ":" not in line:
            continue
        key, value = line.split(":", 1)
        metadata[key.strip()] = value.strip()
    body = text[end + 4 :].lstrip("\r\n")
    return metadata, body


def read_sidecar_metadata(path: Path) -> dict[str, str]:
    candidates = [
        path.with_suffix(path.suffix + ".meta.json"),
        path.with_suffix(".meta.json"),
    ]
    for candidate in candidates:
        if not candidate.exists():
            continue
        try:
            data = json.loads(candidate.read_text(encoding="utf-8"))
        except (json.JSONDecodeError, OSError):
            continue
        if isinstance(data, dict):
            return {str(key): str(value) for key, value in data.items() if value is not None}
    return {}


def discover_documents(raw_dir: Path) -> list[SourceDocument]:
    documents: list[SourceDocument] = []
    for path in sorted(raw_dir.rglob("*")):
        if not path.is_file() or path.name.startswith(".") or path.name.startswith("~$"):
            continue
        suffix = path.suffix.lower()
        if suffix not in SUPPORTED_SUFFIXES:
            continue
        metadata = read_sidecar_metadata(path)
        if suffix == ".md":
            try:
                markdown_metadata, _body = parse_front_matter(read_text_file(path))
                metadata = {**metadata, **markdown_metadata}
            except OSError:
                pass
        category = metadata.get("category") or infer_category(path, raw_dir)
        doc_id = stable_id(str(path))
        documents.append(
            SourceDocument(
                path=path,
                category=category,
                title=metadata.get("title") or infer_title(path),
                suffix=suffix,
                doc_id=doc_id,
                source_url=metadata.get("source_url", ""),
                site=metadata.get("site", ""),
                published_at=metadata.get("published_at", ""),
            )
        )
    return documents


def read_text_file(path: Path) -> str:
    for encoding in ("utf-8", "utf-8-sig", "gb18030"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return path.read_text(encoding="utf-8", errors="ignore")


def read_docx(path: Path) -> str:
    try:
        from docx import Document  # type: ignore
    except ModuleNotFoundError:
        return read_with_textutil(path)

    document = Document(str(path))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def read_doc(path: Path) -> str:
    """Extract text from legacy .doc files (Windows / cross-platform)."""
    # Tier 1: antiword (cross-platform, handles old .doc well)
    antiword = shutil.which("antiword")
    if not antiword:
        # Check common Windows install paths
        for candidate in (
            r"D:\Git\mingw64\bin\antiword.exe",
            r"C:\Program Files\antiword\antiword.exe",
            r"C:\antiword\antiword.exe",
        ):
            if Path(candidate).exists():
                antiword = candidate
                break
    if antiword:
        result = subprocess.run(
            [antiword, str(path)],
            check=False,
            capture_output=True,
            text=True,
        )
        if result.returncode == 0 and result.stdout.strip():
            return result.stdout

    # Tier 2: python-docx can sometimes read .doc (if it's actually docx format)
    try:
        from docx import Document
        doc = Document(str(path))
        parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        if parts:
            return "\n".join(parts)
    except Exception:
        pass

    # Tier 3: textutil (macOS only)
    return read_with_textutil(path)


def read_with_textutil(path: Path) -> str:
    textutil = shutil.which("textutil")
    if not textutil:
        raise RuntimeError("textutil is not available; cannot parse Word documents")

    result = subprocess.run(
        [textutil, "-convert", "txt", "-stdout", str(path)],
        check=False,
        capture_output=True,
        text=True,
    )
    if result.returncode != 0:
        error = result.stderr.strip() or result.stdout.strip()
        raise RuntimeError(f"textutil failed: {error}")
    return result.stdout


def _read_pdf_pymupdf(path: Path) -> str:
    """Extract text using PyMuPDF (fitz) — handles many PDFs better than pypdf."""
    import fitz  # type: ignore[import-untyped]

    doc = fitz.open(str(path))
    pages: list[str] = []
    try:
        for page_index, page in enumerate(doc, start=1):
            page_text = page.get_text("text") or ""
            if page_text.strip():
                pages.append(f"[page {page_index}]\n{page_text}")
    finally:
        doc.close()
    return "\n\n".join(pages)


def _ocr_pdf_pages_pymupdf(path: Path) -> str:
    """Render each page to an image via PyMuPDF, then OCR with pytesseract."""
    try:
        import fitz  # type: ignore[import-untyped]
        import pytesseract  # type: ignore[import-untyped]
        from PIL import Image  # type: ignore[import-untyped]

        # Auto-detect Tesseract if not in PATH
        _tesseract_exe = pytesseract.pytesseract.tesseract_cmd
        if not _tesseract_exe or not Path(_tesseract_exe).exists():
            for _candidate in (
                r"D:\tess\tesseract.exe",
                r"C:\Program Files\Tesseract-OCR\tesseract.exe",
                r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
            ):
                if Path(_candidate).exists():
                    pytesseract.pytesseract.tesseract_cmd = _candidate
                    break
    except ModuleNotFoundError as exc:
        raise RuntimeError(
            f"OCR requires pytesseract + Pillow (and Tesseract installed on the system): {exc}"
        ) from exc

    doc = fitz.open(str(path))
    pages: list[str] = []
    try:
        for page_index, page in enumerate(doc, start=1):
            pix = page.get_pixmap(dpi=200)
            img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
            page_text = pytesseract.image_to_string(img, lang="chi_sim+eng").strip()
            if page_text:
                pages.append(f"[page {page_index}]\n{page_text}")
    finally:
        doc.close()
    return "\n\n".join(pages)


def _read_pdf_pypdf(path: Path) -> str:
    """Extract text with pypdf (legacy fallback)."""
    try:
        from pypdf import PdfReader  # type: ignore[import-untyped]
    except ModuleNotFoundError as exc:
        raise RuntimeError("PDF parser unavailable; install pymupdf or pypdf") from exc

    reader = PdfReader(str(path))
    pages: list[str] = []
    for page_index, page in enumerate(reader.pages, start=1):
        page_text = page.extract_text() or ""
        if page_text.strip():
            pages.append(f"[page {page_index}]\n{page_text}")
    return "\n\n".join(pages)


def _read_pdf_pdftotext(path: Path) -> str:
    """Extract text with pdftotext CLI (last resort)."""
    pdftotext = shutil.which("pdftotext")
    if not pdftotext:
        raise RuntimeError("PDF parser unavailable; install pymupdf, pypdf, or pdftotext")
    result = subprocess.run(
        [pdftotext, "-layout", str(path), "-"],
        check=False,
        capture_output=True,
        text=True,
    )
    if result.returncode != 0:
        error = result.stderr.strip() or result.stdout.strip()
        raise RuntimeError(f"pdftotext failed: {error}")
    return result.stdout


def _text_quality_ok(text: str) -> bool:
    """Check if extracted text has decent quality (not too many garbled chars)."""
    if len(text.strip()) < 100:
        return False
    # U+FFFD = replacement character (garbled encoding)
    garbled = text.count("�")
    total = max(len(text), 1)
    # More than 2% garbled characters = poor quality
    return (garbled / total) < 0.02


def read_pdf(path: Path) -> str:
    """Extract text from a PDF using a tiered strategy:

    1. PyMuPDF (fitz) — best quality, handles mixed text/image PDFs
    2. If text poor quality: OCR via PyMuPDF → pytesseract (for scanned PDFs)
    3. Fall back to pypdf if PyMuPDF unavailable
    4. Last resort: pdftotext CLI
    """

    # ── Tier 1: PyMuPDF ──
    try:
        text = _read_pdf_pymupdf(path)
    except ModuleNotFoundError:
        text = ""
    except Exception:
        text = ""

    if text and _text_quality_ok(text):
        return text

    # ── Tier 2: OCR (scanned / poor quality PDF) ──
    if text and not _text_quality_ok(text):
        garbled = text.count("�")
        print(f"  ⚠ PyMuPDF: {len(text.strip())} chars ({garbled} garbled) from {path.name}, trying OCR...")
    elif not text or len(text.strip()) < 100:
        print(f"  ⚠ PyMuPDF extracted only {len(text.strip()) if text else 0} chars from {path.name}, trying OCR...")

    try:
        ocr_text = _ocr_pdf_pages_pymupdf(path)
        if ocr_text and _text_quality_ok(ocr_text):
            print(f"  ✓ OCR succeeded: {len(ocr_text.strip())} chars from {path.name}")
            return ocr_text
    except (ModuleNotFoundError, RuntimeError) as exc:
        print(f"  ⚠ OCR unavailable for {path.name}: {exc}")
    except Exception as exc:
        print(f"  ⚠ OCR failed for {path.name}: {exc}")

    if text.strip():
        return text

    # ── Tier 3: pypdf ──
    try:
        text = _read_pdf_pypdf(path)
        if text and len(text.strip()) >= 100:
            return text
    except (ModuleNotFoundError, RuntimeError):
        pass
    except Exception:
        pass

    # ── Tier 4: pdftotext ──
    return _read_pdf_pdftotext(path)


def extract_text(document: SourceDocument) -> str:
    if document.suffix in {".txt", ".md"}:
        metadata, body = parse_front_matter(read_text_file(document.path))
        return body if metadata else read_text_file(document.path)
    if document.suffix == ".docx":
        return read_docx(document.path)
    if document.suffix == ".doc":
        return read_doc(document.path)
    if document.suffix == ".pdf":
        return read_pdf(document.path)
    raise RuntimeError(f"unsupported file type: {document.suffix}")


def is_heading(line: str) -> bool:
    line = line.strip()
    if not line or len(line) > 60:
        return False
    patterns = [
        r"^[一二三四五六七八九十]+[、.．]",
        r"^第[一二三四五六七八九十0-9]+[章节部分条]",
        r"^\d+(\.\d+)*[、.．]\s*",
        r".*(课程说明|课程简介|课程目标|毕业要求|课程体系|课程考核|教学内容|培养目标)$",
    ]
    return any(re.match(pattern, line) for pattern in patterns)


def paragraph_blocks(text: str) -> Iterable[tuple[str, str]]:
    current_heading = ""
    buffer: list[str] = []

    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            if buffer:
                yield current_heading, "\n".join(buffer).strip()
                buffer = []
            continue

        if is_heading(line):
            if buffer:
                yield current_heading, "\n".join(buffer).strip()
                buffer = []
            current_heading = line
            buffer.append(line)
        else:
            buffer.append(line)

    if buffer:
        yield current_heading, "\n".join(buffer).strip()


def split_long_block(text: str, max_chars: int, overlap: int) -> list[str]:
    if len(text) <= max_chars:
        return [text]

    parts: list[str] = []
    start = 0
    while start < len(text):
        end = min(start + max_chars, len(text))
        candidate = text[start:end]
        if end < len(text):
            split_at = max(candidate.rfind("\n"), candidate.rfind("。"), candidate.rfind("；"))
            if split_at > max_chars * 0.5:
                end = start + split_at + 1
                candidate = text[start:end]
        parts.append(candidate.strip())
        if end >= len(text):
            break
        next_start = max(0, end - overlap)
        boundary_window = text[next_start : min(end + 1, len(text))]
        boundary_offsets = [
            offset + 1
            for offset, char in enumerate(boundary_window)
            if char in {"\n", "。", "；"}
        ]
        start = next_start + boundary_offsets[0] if boundary_offsets else next_start
    return [part for part in parts if part]


def chunk_text(
    document: SourceDocument,
    text: str,
    max_chars: int,
    overlap: int,
) -> list[Chunk]:
    chunks: list[Chunk] = []
    pending_heading = ""
    pending_parts: list[str] = []
    pending_size = 0
    # Carry over the last paragraph from previous flush as overlap
    overlap_carry: str | None = None

    def flush(keep_last: bool = False) -> str | None:
        nonlocal pending_heading, pending_parts, pending_size, overlap_carry
        if not pending_parts:
            return None
        # Prepend overlap from previous chunk
        if overlap_carry:
            pending_parts.insert(0, overlap_carry)
            pending_size += len(overlap_carry)
        combined = "\n\n".join(pending_parts).strip()
        for part in split_long_block(combined, max_chars, overlap):
            index = len(chunks)
            chunk_key = f"{document.doc_id}:{index}:{part[:80]}"
            chunks.append(
                Chunk(
                    chunk_id=stable_id(chunk_key),
                    doc_id=document.doc_id,
                    title=document.title,
                    category=document.category,
                    source_file=str(document.path),
                    chunk_index=index,
                    heading=pending_heading,
                    text=part,
                    source_url=document.source_url,
                    site=document.site,
                    published_at=document.published_at,
                )
            )
        # Carry last paragraph as overlap for next chunk
        if keep_last and pending_parts and len(pending_parts) > 1:
            overlap_carry = pending_parts[-1]
        else:
            overlap_carry = None
        pending_heading = ""
        pending_parts = []
        pending_size = 0
        return overlap_carry

    for heading, block in paragraph_blocks(text):
        if not block:
            continue
        block_size = len(block)
        if pending_parts and pending_size + block_size > max_chars:
            flush(keep_last=True)
        if not pending_heading and heading:
            pending_heading = heading
        pending_parts.append(block)
        pending_size += block_size

    flush()
    return chunks


def write_jsonl(chunks: list[Chunk], output_file: Path) -> None:
    output_file.parent.mkdir(parents=True, exist_ok=True)
    with output_file.open("w", encoding="utf-8") as file:
        for chunk in chunks:
            file.write(json.dumps(asdict(chunk), ensure_ascii=False) + "\n")


def write_report(report: dict, report_file: Path) -> None:
    report_file.parent.mkdir(parents=True, exist_ok=True)
    report_file.write_text(
        json.dumps(report, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )


def build_chunks(args: argparse.Namespace) -> int:
    documents = discover_documents(args.raw_dir)
    all_chunks: list[Chunk] = []
    parsed: list[dict] = []
    failed: list[dict] = []

    for document in documents:
        try:
            text = normalize_text(extract_text(document))
            if len(text) < args.min_chars:
                raise RuntimeError(f"extracted text too short: {len(text)} chars")
            chunks = chunk_text(document, text, args.max_chars, args.overlap)
            all_chunks.extend(chunks)
            parsed.append(
                {
                    "source_file": str(document.path),
                    "category": document.category,
                    "title": document.title,
                    "text_chars": len(text),
                    "chunks": len(chunks),
                    "source_url": document.source_url,
                    "site": document.site,
                    "published_at": document.published_at,
                }
            )
        except Exception as exc:
            failed.append(
                {
                    "source_file": str(document.path),
                    "category": document.category,
                    "title": document.title,
                    "error": str(exc),
                }
            )

    write_jsonl(all_chunks, args.output)
    report = {
        "built_at": datetime.now(timezone.utc).isoformat(),
        "raw_dir": str(args.raw_dir),
        "output": str(args.output),
        "document_count": len(documents),
        "parsed_count": len(parsed),
        "failed_count": len(failed),
        "chunk_count": len(all_chunks),
        "parsed": parsed,
        "failed": failed,
    }
    write_report(report, args.report)

    print(f"Documents: {len(documents)}")
    print(f"Parsed: {len(parsed)}")
    print(f"Failed: {len(failed)}")
    print(f"Chunks: {len(all_chunks)}")
    print(f"Output: {args.output}")
    print(f"Report: {args.report}")
    if failed:
        print("Failed files:")
        for item in failed:
            print(f"- {item['source_file']}: {item['error']}")
    return 0 if all_chunks else 1


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--raw-dir", type=Path, default=RAW_DIR)
    parser.add_argument("--output", type=Path, default=OUTPUT_FILE)
    parser.add_argument("--report", type=Path, default=REPORT_FILE)
    parser.add_argument("--max-chars", type=int, default=1200)
    parser.add_argument("--overlap", type=int, default=120)
    parser.add_argument("--min-chars", type=int, default=80)
    return parser.parse_args()


if __name__ == "__main__":
    raise SystemExit(build_chunks(parse_args()))
